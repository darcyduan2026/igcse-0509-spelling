from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import html


WORKDIR = Path(".")
HTML_FILES = [
    Path("[唐代] 杜甫 《严郑公宅同咏竹》 批注海报.html"),
    Path("[宋代] 苏轼 《於潜僧绿筠轩》 批注海报.html"),
    Path("[明代] 刘基 《蜀贾》 批注海报.html"),
    Path("[明代] 宋应星 《天工开物》序 批注海报.html"),
    Path("[明代] 宋濂 《记李歌》 批注海报.html"),
    Path("[明代] 张岱 《湖心亭看雪》 批注海报.html"),
    Path("yangming_annotated_poster.html"),
]

BLUE = RGBColor(29, 95, 191)
RED = RGBColor(198, 61, 50)
GOLD = RGBColor(184, 121, 26)
INK = RGBColor(34, 40, 49)
MUTED = RGBColor(109, 103, 95)


def text_content(el) -> str:
    return " ".join("".join(el.itertext()).split())


def out_path_for(html_path: Path, title: str) -> Path:
    if html_path.name == "yangming_annotated_poster.html":
        return Path("阳明子_四民异业而同道_批注海报.docx")
    return Path(title).with_suffix(".docx")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def has_class(el, class_name: str) -> bool:
    return class_name in (el.get("class") or "").split()


def by_class(tree, tag: str, class_name: str):
    return tree.xpath(f".//{tag}[contains(concat(' ', normalize-space(@class), ' '), ' {class_name} ')]")


def style_document(doc: Document):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Songti SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size in (("Heading 1", 18), ("Heading 2", 14)):
        style = styles[name]
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = INK
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_title(doc: Document, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.size = Pt(20)
    run.font.color.rgb = INK

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    run = p2.add_run("蓝色：拼音    红色：字义/句意    金色下划线：人物、地名（Word批注）")
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED


def add_core_vocab(doc: Document, tree):
    cards = by_class(tree, "*", "core-card")
    if not cards:
        return
    doc.add_heading("核心词汇", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "词语"
    hdr[1].text = "提示"
    for cell in hdr:
        set_cell_shading(cell, "F2E7D2")
        set_cell_margins(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "PingFang SC"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    for card in cards:
        row = table.add_row().cells
        strong = card.xpath(".//strong")
        span = card.xpath(".//span")
        row[0].text = text_content(strong[0]) if strong else ""
        row[1].text = text_content(span[0]) if span else ""
        for cell in row:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def collect_comments(tree) -> dict[str, str]:
    notes = {}
    for aside in by_class(tree, "aside", "note"):
        target = aside.get("data-target")
        if not target:
            continue
        title_el = aside.xpath(".//h2")
        p_el = aside.xpath(".//p")
        small_el = aside.xpath(".//small")
        parts = []
        if title_el:
            parts.append(text_content(title_el[0]))
        if p_el:
            parts.append(text_content(p_el[0]))
        if small_el:
            parts.append(text_content(small_el[0]))
        notes[target] = "\n".join(part for part in parts if part)
    return notes


def style_run(run, cls: str | None, is_comment_target=False):
    run.font.name = "Songti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    run.font.size = Pt(12)
    if cls and "py" in cls.split():
        run.font.name = "PingFang SC"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        run.font.color.rgb = BLUE
        run.bold = True
        run.font.size = Pt(10.5)
    elif cls and "meaning" in cls.split():
        run.font.name = "PingFang SC"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        run.font.color.rgb = RED
        run.bold = True
        run.font.size = Pt(10.5)
    elif cls and ("person" in cls.split() or "highlight" in cls.split()):
        run.font.color.rgb = INK
        run.font.underline = True
        run.bold = True
    if is_comment_target:
        run.font.color.rgb = GOLD
        run.font.underline = True
        run.bold = True


def add_text_run(paragraph, text: str, cls: str | None = None, comment_text: str | None = None):
    if not text:
        return
    run = paragraph.add_run(text)
    style_run(run, cls, bool(comment_text))
    if comment_text:
        paragraph.part.document.add_comment(run, text=comment_text, author="人物/地名批注", initials="注")


def add_node_content(paragraph, node, comments: dict[str, str], seen_comment_targets: set[str]):
    if node.text:
        cls = node.get("class")
        target_id = node.get("id")
        comment_text = comments.get(target_id)
        add_text_run(paragraph, node.text, cls, comment_text)
        if comment_text and target_id:
            seen_comment_targets.add(target_id)
    for child in node:
        add_node_content(paragraph, child, comments, seen_comment_targets)
        if child.tail:
            add_text_run(paragraph, child.tail)


def add_main_text(doc: Document, tree, comments) -> set[str]:
    seen_comment_targets: set[str] = set()
    doc.add_heading("正文批注", level=2)
    main_sections = by_class(tree, "section", "main-text")
    html_ps = main_sections[0].xpath(".//p") if main_sections else []
    for html_p in html_ps:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.space_after = Pt(8)
        add_node_content(p, html_p, comments, seen_comment_targets)
    return seen_comment_targets


def add_orphan_comments(doc: Document, comments: dict[str, str], seen_comment_targets: set[str]):
    missing = [(target, text) for target, text in comments.items() if target not in seen_comment_targets]
    if not missing:
        return
    doc.add_heading("人物 / 地名批注", level=2)
    p = doc.add_paragraph()
    for index, (_target, comment_text) in enumerate(missing):
        label = comment_text.splitlines()[0]
        if index:
            add_text_run(p, "　")
        add_text_run(p, label, "person", comment_text)


def add_footer(doc: Document, tree):
    footers = by_class(tree, "footer", "footer")
    spans = footers[0].xpath(".//span") if footers else []
    footer_texts = [text_content(span) for span in spans]
    if not footer_texts:
        return
    doc.add_heading("自学提示", level=2)
    for text in footer_texts:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(12)
        run = p.add_run(text)
        run.font.name = "PingFang SC"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        run.font.size = Pt(10.5)
        run.font.color.rgb = MUTED


def build_docx(html_path: Path):
    raw = html_path.read_text(encoding="utf-8")
    tree = html.fromstring(raw)
    h1 = tree.xpath(".//h1")
    title = text_content(h1[0]) if h1 else html_path.stem
    out_path = out_path_for(html_path, title)

    doc = Document()
    style_document(doc)
    add_title(doc, title)
    add_core_vocab(doc, tree)
    comments = collect_comments(tree)
    seen_comment_targets = add_main_text(doc, tree, comments)
    add_orphan_comments(doc, comments, seen_comment_targets)
    add_footer(doc, tree)
    doc.save(out_path)
    return out_path


for html_path in HTML_FILES:
    print(build_docx(html_path))
